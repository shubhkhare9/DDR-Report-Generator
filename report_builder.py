"""
report_builder.py — Builds DDR as Word (.docx) and PDF (.pdf)
Uses python-docx for Word, reportlab for PDF
Images from source documents are embedded into reports
"""
import os
import io
import tempfile
from datetime import datetime
from typing import Dict, List

# ── Word ──────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PDF ───────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _severity_color(severity: str):
    s = severity.strip().lower()
    if "high" in s:
        return RGBColor(0xC0, 0x00, 0x00)   # dark red
    if "medium" in s:
        return RGBColor(0xFF, 0x8C, 0x00)   # orange
    return RGBColor(0x00, 0x70, 0x00)       # green


def _priority_color(priority: str):
    p = priority.strip().lower()
    if "immediate" in p:
        return colors.HexColor("#C00000")
    if "short" in p:
        return colors.HexColor("#FF8C00")
    return colors.HexColor("#007000")


def _pick_images_for_area(area_name: str, inspection_data: Dict, thermal_data: Dict):
    """
    Heuristic: return up to 2 inspection + 1 thermal image bytes for an area.
    Falls back to page-order selection when area name can't be matched.
    """
    area_lower = area_name.lower()

    # map area keywords → approx page ranges in inspection PDF
    area_page_map = {
        "hall": [3],
        "bedroom": [3, 4],
        "master bedroom": [4, 5],
        "kitchen": [4],
        "parking": [5, 6],
        "common bathroom": [3, 6, 7],
        "external wall": [5],
    }

    target_pages = []
    for kw, pages in area_page_map.items():
        if kw in area_lower:
            target_pages = pages
            break

    all_insp = inspection_data.get("images", [])
    if target_pages:
        insp_imgs = [img for img in all_insp if img["page"] in target_pages][:2]
    else:
        insp_imgs = all_insp[:2]

    therm_imgs = thermal_data.get("images", [])[:1]

    return insp_imgs, therm_imgs


# ─────────────────────────────────────────────────────────────────────────────
# WORD REPORT
# ─────────────────────────────────────────────────────────────────────────────

def build_word_report(ddr_content: Dict, inspection_data: Dict, thermal_data: Dict) -> str:
    """Build a formatted .docx DDR report. Returns path to the saved file."""

    doc = Document()
    prop = ddr_content.get("property_info", {})
    sections = ddr_content.get("sections", [])

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover heading ─────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DETAILED DIAGNOSTIC REPORT (DDR)")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Waterproofing & Structural Inspection Analysis")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()  # spacer

    # ── Property Info Table ───────────────────────────────────
    info_rows = [
        ("Property Type", prop.get("property_type", "Not Available")),
        ("Flat / Unit", prop.get("flat_number", "Not Available")),
        ("Number of Floors", prop.get("floors", "Not Available")),
        ("Inspection Date", prop.get("inspection_date", "Not Available")),
        ("Inspected By", prop.get("inspected_by", "Not Available")),
        ("Inspection Score", prop.get("inspection_score", "Not Available")),
        ("Previous Structural Audit", prop.get("previous_audit", "Not Available")),
        ("Previous Repair Work", prop.get("previous_repair", "Not Available")),
        ("Report Generated On", datetime.now().strftime("%d %b %Y")),
    ]

    tbl = doc.add_table(rows=len(info_rows), cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(10)

    for i, (label, value) in enumerate(info_rows):
        row = tbl.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]

        # Header cell shading
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "D9E1F2")
        shading_elm.set(qn("w:val"), "clear")
        cell_label._tc.get_or_add_tcPr().append(shading_elm)

        lbl_run = cell_label.paragraphs[0].add_run(label)
        lbl_run.bold = True
        lbl_run.font.size = Pt(10)

        val_run = cell_value.paragraphs[0].add_run(str(value))
        val_run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Seven DDR Sections ────────────────────────────────────
    for sec in sections:
        heading = doc.add_heading(sec.get("title", ""), level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        body = doc.add_paragraph(sec.get("content", ""))
        body.paragraph_format.space_after = Pt(6)

        # ── Section 2: Area-wise Observations ────────────────
        if sec.get("id") == "s2":
            for area in sec.get("areas", []):
                area_heading = doc.add_heading(f"  ► {area.get('area_name', '')}", level=2)
                area_heading.runs[0].font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

                details = [
                    ("Problem Observed (Negative Side)", area.get("negative_side", "Not Available")),
                    ("Source of Issue (Positive Side)", area.get("positive_side", "Not Available")),
                    ("Thermal Analysis Finding", area.get("thermal_finding", "Not Available")),
                ]
                for label, val in details:
                    p = doc.add_paragraph(style="List Bullet")
                    run_lbl = p.add_run(f"{label}: ")
                    run_lbl.bold = True
                    run_lbl.font.size = Pt(10)
                    run_val = p.add_run(val)
                    run_val.font.size = Pt(10)

                # Add images for this area
                insp_imgs, therm_imgs = _pick_images_for_area(
                    area.get("area_name", ""), inspection_data, thermal_data
                )

                if insp_imgs or therm_imgs:
                    doc.add_paragraph("Supporting Images:").runs[0].bold = True
                    img_count = 0
                    for img in insp_imgs:
                        try:
                            img_stream = io.BytesIO(img["bytes"])
                            doc.add_picture(img_stream, width=Cm(7))
                            cap = doc.add_paragraph(
                                f"Inspection Photo — Page {img['page']}"
                            )
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].font.size = Pt(8)
                            cap.runs[0].italic = True
                            img_count += 1
                        except Exception:
                            doc.add_paragraph("[Inspection image could not be embedded]")

                    for img in therm_imgs:
                        try:
                            img_stream = io.BytesIO(img["bytes"])
                            doc.add_picture(img_stream, width=Cm(7))
                            cap = doc.add_paragraph(
                                f"Thermal Image — Page {img['page']}"
                            )
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].font.size = Pt(8)
                            cap.runs[0].italic = True
                        except Exception:
                            doc.add_paragraph("[Thermal image could not be embedded]")

                doc.add_paragraph()  # spacer after each area

        # ── Section 4: Severity Table ─────────────────────────
        elif sec.get("id") == "s4" and sec.get("severity_table"):
            rows = sec["severity_table"]
            tbl = doc.add_table(rows=len(rows) + 1, cols=3)
            tbl.style = "Table Grid"
            headers = ["Area / Location", "Severity Level", "Reasoning"]
            header_row = tbl.rows[0]
            for j, h in enumerate(headers):
                cell = header_row.cells[j]
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1F497D")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

            for i, row_data in enumerate(rows):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(
                    row_data.get("area", "")
                ).font.size = Pt(10)

                sev_text = row_data.get("severity", "")
                sev_run = row.cells[1].paragraphs[0].add_run(sev_text)
                sev_run.bold = True
                sev_run.font.size = Pt(10)
                sev_run.font.color.rgb = _severity_color(sev_text)

                row.cells[2].paragraphs[0].add_run(
                    row_data.get("reasoning", "")
                ).font.size = Pt(10)

            doc.add_paragraph()

        # ── Section 5: Actions Table ──────────────────────────
        elif sec.get("id") == "s5" and sec.get("actions"):
            actions = sec["actions"]
            tbl = doc.add_table(rows=len(actions) + 1, cols=3)
            tbl.style = "Table Grid"
            headers = ["Area / Location", "Recommended Action", "Priority"]
            header_row = tbl.rows[0]
            for j, h in enumerate(headers):
                cell = header_row.cells[j]
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "375623")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

            for i, action in enumerate(actions):
                row = tbl.rows[i + 1]
                row.cells[0].paragraphs[0].add_run(
                    action.get("area", "")
                ).font.size = Pt(10)
                row.cells[1].paragraphs[0].add_run(
                    action.get("action", "")
                ).font.size = Pt(10)
                prio_text = action.get("priority", "")
                prio_run = row.cells[2].paragraphs[0].add_run(prio_text)
                prio_run.bold = True
                prio_run.font.size = Pt(10)
                prio_run.font.color.rgb = _severity_color(prio_text)

            doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "This report was generated using an AI-powered DDR system. "
        "All findings are based solely on the provided inspection and thermal documents."
    )
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save
    out_path = os.path.join(tempfile.gettempdir(), "DDR_Report.docx")
    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# PDF REPORT
# ─────────────────────────────────────────────────────────────────────────────

def build_pdf_report(ddr_content: Dict, inspection_data: Dict, thermal_data: Dict) -> str:
    """Build a formatted .pdf DDR report. Returns path to the saved file."""

    out_path = os.path.join(tempfile.gettempdir(), "DDR_Report.pdf")
    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "DDRTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#1F497D"),
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "DDRSubtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#707070"),
        spaceAfter=16,
        alignment=TA_CENTER,
    )
    h1_style = ParagraphStyle(
        "DDRH1",
        parent=styles["Heading1"],
        fontSize=13,
        textColor=colors.HexColor("#1F497D"),
        spaceBefore=14,
        spaceAfter=6,
        borderPad=3,
    )
    h2_style = ParagraphStyle(
        "DDRH2",
        parent=styles["Heading2"],
        fontSize=11,
        textColor=colors.HexColor("#C05000"),
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "DDRBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
    )
    label_style = ParagraphStyle(
        "DDRLabel",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#333333"),
    )
    caption_style = ParagraphStyle(
        "DDRCaption",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    footer_style = ParagraphStyle(
        "DDRFooter",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceBefore=20,
    )

    story = []
    prop = ddr_content.get("property_info", {})
    sections = ddr_content.get("sections", [])

    # ── Cover ─────────────────────────────────────────────────
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("DETAILED DIAGNOSTIC REPORT (DDR)", title_style))
    story.append(Paragraph("Waterproofing & Structural Inspection Analysis", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1F497D")))
    story.append(Spacer(1, 0.4 * cm))

    # ── Property Info Table ───────────────────────────────────
    info_data = [
        [Paragraph("<b>Field</b>", label_style), Paragraph("<b>Details</b>", label_style)],
        ["Property Type", prop.get("property_type", "Not Available")],
        ["Flat / Unit", prop.get("flat_number", "Not Available")],
        ["Number of Floors", prop.get("floors", "Not Available")],
        ["Inspection Date", prop.get("inspection_date", "Not Available")],
        ["Inspected By", prop.get("inspected_by", "Not Available")],
        ["Inspection Score", prop.get("inspection_score", "Not Available")],
        ["Previous Structural Audit", prop.get("previous_audit", "Not Available")],
        ["Previous Repair Work", prop.get("previous_repair", "Not Available")],
        ["Report Generated On", datetime.now().strftime("%d %b %Y")],
    ]

    info_table = Table(info_data, colWidths=[6 * cm, 10.5 * cm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F497D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#D9E1F2")),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (1, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FB")]),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D0D0D0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.6 * cm))

    # ── Seven DDR Sections ────────────────────────────────────
    for sec in sections:
        story.append(Paragraph(sec.get("title", ""), h1_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D0D0D0")))
        story.append(Spacer(1, 0.2 * cm))

        content_text = sec.get("content", "")
        if content_text:
            story.append(Paragraph(content_text, body_style))

        # ── Section 2: Area-wise Observations ────────────────
        if sec.get("id") == "s2":
            for area in sec.get("areas", []):
                story.append(Paragraph(f"► {area.get('area_name', '')}", h2_style))

                area_data = [
                    [
                        Paragraph("<b>Problem Observed (Negative Side)</b>", label_style),
                        Paragraph(area.get("negative_side", "Not Available"), body_style),
                    ],
                    [
                        Paragraph("<b>Source of Issue (Positive Side)</b>", label_style),
                        Paragraph(area.get("positive_side", "Not Available"), body_style),
                    ],
                    [
                        Paragraph("<b>Thermal Analysis Finding</b>", label_style),
                        Paragraph(area.get("thermal_finding", "Not Available"), body_style),
                    ],
                ]
                area_tbl = Table(area_data, colWidths=[5 * cm, 11.5 * cm])
                area_tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#FFF3E0")),
                    ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#C0C0C0")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E0E0E0")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ]))
                story.append(area_tbl)
                story.append(Spacer(1, 0.2 * cm))

                # Images
                insp_imgs, therm_imgs = _pick_images_for_area(
                    area.get("area_name", ""), inspection_data, thermal_data
                )
                img_elements = []
                for img in insp_imgs:
                    try:
                        img_stream = io.BytesIO(img["bytes"])
                        rl_img = RLImage(img_stream, width=6 * cm, height=4.5 * cm)
                        img_elements.append(rl_img)
                        img_elements.append(
                            Paragraph(f"Inspection Photo — Page {img['page']}", caption_style)
                        )
                    except Exception:
                        pass

                for img in therm_imgs:
                    try:
                        img_stream = io.BytesIO(img["bytes"])
                        rl_img = RLImage(img_stream, width=6 * cm, height=4.5 * cm)
                        img_elements.append(rl_img)
                        img_elements.append(
                            Paragraph(f"Thermal Image — Page {img['page']}", caption_style)
                        )
                    except Exception:
                        pass

                if img_elements:
                    story.extend(img_elements)

                story.append(Spacer(1, 0.4 * cm))

        # ── Section 4: Severity Table ─────────────────────────
        elif sec.get("id") == "s4" and sec.get("severity_table"):
            rows = sec["severity_table"]
            sev_data = [[
                Paragraph("<b>Area / Location</b>", label_style),
                Paragraph("<b>Severity</b>", label_style),
                Paragraph("<b>Reasoning</b>", label_style),
            ]]
            for r in rows:
                sev_text = r.get("severity", "")
                color_map = {"high": "#C00000", "medium": "#FF8C00", "low": "#007000"}
                hex_color = "#333333"
                for k, v in color_map.items():
                    if k in sev_text.lower():
                        hex_color = v
                        break

                sev_data.append([
                    Paragraph(r.get("area", ""), body_style),
                    Paragraph(
                        f'<font color="{hex_color}"><b>{sev_text}</b></font>', body_style
                    ),
                    Paragraph(r.get("reasoning", ""), body_style),
                ])

            sev_tbl = Table(sev_data, colWidths=[5 * cm, 2.5 * cm, 9 * cm])
            sev_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F497D")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FB")]),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D0D0D0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(sev_tbl)
            story.append(Spacer(1, 0.4 * cm))

        # ── Section 5: Recommended Actions ───────────────────
        elif sec.get("id") == "s5" and sec.get("actions"):
            actions = sec["actions"]
            act_data = [[
                Paragraph("<b>Area</b>", label_style),
                Paragraph("<b>Recommended Action</b>", label_style),
                Paragraph("<b>Priority</b>", label_style),
            ]]
            for a in actions:
                prio_text = a.get("priority", "")
                color_map = {
                    "immediate": "#C00000",
                    "short": "#FF8C00",
                    "long": "#007000",
                }
                hex_color = "#333333"
                for k, v in color_map.items():
                    if k in prio_text.lower():
                        hex_color = v
                        break

                act_data.append([
                    Paragraph(a.get("area", ""), body_style),
                    Paragraph(a.get("action", ""), body_style),
                    Paragraph(
                        f'<font color="{hex_color}"><b>{prio_text}</b></font>', body_style
                    ),
                ])

            act_tbl = Table(act_data, colWidths=[4 * cm, 9 * cm, 3.5 * cm])
            act_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#375623")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F5F0")]),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D0D0D0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(act_tbl)
            story.append(Spacer(1, 0.4 * cm))

    # ── Footer ────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#C0C0C0")))
    story.append(
        Paragraph(
            "This report was generated using an AI-powered DDR system. "
            "All findings are based solely on the provided inspection and thermal documents.",
            footer_style,
        )
    )

    doc.build(story)
    return out_path
